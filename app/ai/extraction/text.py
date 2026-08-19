"""Text extraction from uploaded documents.

Every failure here produces a `DocumentParseError` whose message is written for
a Glimmora user, not a developer — the record is still creatable by hand
(master brief section 24).
"""

from __future__ import annotations

import io
import re

from app.core.errors import DocumentParseError, UnsupportedMediaTypeError
from app.core.logging import get_logger

logger = get_logger("ai.text")

#: Extension -> the magic bytes a file of that type must actually start with.
#: The extension alone is never trusted (SECURITY.md section 6).
MAGIC_BYTES: dict[str, tuple[bytes, ...]] = {
    "pdf": (b"%PDF-",),
    "docx": (b"PK\x03\x04",),
    "doc": (b"\xd0\xcf\x11\xe0",),
    "txt": (),
    "md": (),
}

SUPPORTED_EXTENSIONS = tuple(MAGIC_BYTES)

#: Below this, a "document" is almost certainly a scan or an empty template.
MIN_USEFUL_CHARACTERS = 80


def extension_of(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def validate_document(filename: str, content: bytes, *, max_bytes: int) -> str:
    """Return the validated extension, or raise a user-facing error."""
    extension = extension_of(filename)

    if extension not in MAGIC_BYTES:
        raise UnsupportedMediaTypeError(
            "That file type is not supported. Upload a PDF, Word document or plain text file."
        )
    if not content:
        raise DocumentParseError("That file is empty.")
    if len(content) > max_bytes:
        raise DocumentParseError(
            f"That file is larger than the {max_bytes // (1024 * 1024)} MB limit."
        )

    signatures = MAGIC_BYTES[extension]
    if signatures and not any(content.startswith(signature) for signature in signatures):
        raise DocumentParseError(
            "That file does not look like the type its name suggests. "
            "Re-save it and try again, or paste the text instead."
        )
    return extension


def normalise_text(raw: str) -> str:
    """Collapse the whitespace noise that PDF extraction leaves behind."""
    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace(" ", " ").replace("•", "- ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    return text.strip()


def extract_text(filename: str, content: bytes, *, max_bytes: int) -> str:
    """Extract readable text from an uploaded document."""
    extension = validate_document(filename, content, max_bytes=max_bytes)

    if extension in {"txt", "md"}:
        text = _decode(content)
    elif extension == "pdf":
        text = _extract_pdf(content)
    elif extension == "docx":
        text = _extract_docx(content)
    else:  # legacy .doc
        raise DocumentParseError(
            "Legacy .doc files are not supported. Save it as .docx or PDF, "
            "or paste the text instead."
        )

    text = normalise_text(text)
    if len(text) < MIN_USEFUL_CHARACTERS:
        raise DocumentParseError(
            "Unable to read enough text from this document. It may be a scanned image. "
            "You can paste the text or enter the details manually."
        )
    return text


def _decode(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    raise DocumentParseError("Unable to read the text encoding of that file.")


def _extract_pdf(content: bytes) -> str:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise DocumentParseError("PDF reading is not available on this server.") from exc

    try:
        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            # A JD is a few pages; a 200-page document is not one, and reading it
            # all would stall the request.
            for page in pdf.pages[:20]:
                pages.append(page.extract_text() or "")
        return "\n\n".join(pages)
    except DocumentParseError:
        raise
    except Exception as exc:
        logger.warning("pdf_extraction_failed", error=str(exc))
        raise DocumentParseError() from exc


def _extract_docx(content: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise DocumentParseError("Word reading is not available on this server.") from exc

    try:
        document = docx.Document(io.BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs]
        # Rate cards and skill matrices are routinely in tables, not paragraphs.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except DocumentParseError:
        raise
    except Exception as exc:
        logger.warning("docx_extraction_failed", error=str(exc))
        raise DocumentParseError() from exc


__all__ = [
    "MAGIC_BYTES",
    "MIN_USEFUL_CHARACTERS",
    "SUPPORTED_EXTENSIONS",
    "extension_of",
    "extract_text",
    "normalise_text",
    "validate_document",
]
